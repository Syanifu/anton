import json
import sys
import os
import argparse
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from docx import Document

def create_pdf(content, output_path):
    """
    Generates a PDF file from the provided content.
    """
    doc = SimpleDocTemplate(output_path, pagesize=letter)
    
    styles = getSampleStyleSheet()
    story = []

    # Title
    if 'title' in content:
        story.append(Paragraph(content['title'], styles['Title']))
        story.append(Spacer(1, 12))

    # Sections
    if 'sections' in content:
        for section in content['sections']:
            if 'heading' in section:
                story.append(Paragraph(section['heading'], styles['Heading2']))
                story.append(Spacer(1, 6))
            
            if 'body' in section:
                # Handle newlines in body text by splitting or using appropriate style
                # For simplicity, we just add it as a paragraph
                story.append(Paragraph(section['body'], styles['Normal']))
                story.append(Spacer(1, 12))

    doc.build(story)
    print(f"PDF generated at: {output_path}")

def create_docx(content, output_path):
    """
    Generates a DOCX file from the provided content.
    """
    doc = Document()

    # Title
    if 'title' in content:
        doc.add_heading(content['title'], 0)

    # Sections
    if 'sections' in content:
        for section in content['sections']:
            if 'heading' in section:
                doc.add_heading(section['heading'], level=2)
            
            if 'body' in section:
                doc.add_paragraph(section['body'])

    doc.save(output_path)
    print(f"DOCX generated at: {output_path}")

def main():
    parser = argparse.ArgumentParser(description="Generate PDF or DOCX documents from JSON content.")
    parser.add_argument("--content", help="JSON string containing document content")
    parser.add_argument("--content-file", help="Path to JSON file containing document content")
    parser.add_argument("--output", required=True, help="Path to save the generated file")
    parser.add_argument("--format", choices=['pdf', 'docx'], required=True, help="Output format (pdf or docx)")

    args = parser.parse_args()

    try:
        if args.content_file:
            with open(args.content_file, 'r') as f:
                content = json.load(f)
        elif args.content:
            content = json.loads(args.content)
        else:
            print("Error: Either --content or --content-file must be provided", file=sys.stderr)
            sys.exit(1)
    except json.JSONDecodeError as e:
        print(f"Error parsing JSON content: {e}", file=sys.stderr)
        sys.exit(1)

    output_path = args.output
    
    # Ensure output directory exists
    output_dir = os.path.dirname(output_path)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)

    if args.format == 'pdf':
        create_pdf(content, output_path)
    elif args.format == 'docx':
        create_docx(content, output_path)

if __name__ == "__main__":
    main()
